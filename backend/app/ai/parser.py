"""CV Text Extraction Module - Extracts raw text from PDF and DOCX files."""

import pdfplumber
from docx import Document


def extract_text(file_path: str, file_type: str) -> str:
    """Extract text from a CV file (PDF or DOCX)."""
    if file_type == "pdf":
        return _extract_from_pdf(file_path)
    elif file_type == "docx":
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)
